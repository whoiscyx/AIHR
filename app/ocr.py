"""简历文本提取（混合策略）。

OCR 引擎优先级：
  1. 本机已安装的 Umi-OCR（PaddleOCR-json.exe）—— 底层即 PaddleOCR，中文精度高、已装即用。
     通过管道（stdin/stdout 收发 JSON）调用，作为子进程懒加载，不依赖 Umi-OCR 自带 Python 环境。
  2. 若找不到 Umi-OCR 可执行文件，则自动回退到 EasyOCR。

提取策略：
- PDF：先用 PyMuPDF 抽取文本；若文本量过小（扫描件/图片型 PDF），则把每页渲染成图片再走 OCR。
- 图片（png/jpg/jpeg/bmp/webp/tif/tiff）：直接 OCR。
- DOCX：python-docx 抽取段落与表格文本。

Umi-OCR 引擎路径可通过环境变量 UMI_OCR_EXE 覆盖；未设置时使用下方默认路径。
"""

import os
import sys
import json
import time
import atexit
import base64
import subprocess
import threading

from pathlib import Path

# ── Umi-OCR（PaddleOCR-json）引擎路径 ────────────────────────────────────────
# 默认指向本机安装的 Umi-OCR Paddle 版自带的 PaddleOCR-json 引擎。
# 如需换版本/路径，设置环境变量 UMI_OCR_EXE 即可（指向 PaddleOCR-json.exe）。
DEFAULT_UMI_OCR_EXE = os.environ.get(
    "UMI_OCR_EXE",
    r"D:\software\OCR\Umi-OCR_Paddle_v2.1.5\UmiOCR-data\plugins\win7_x64_PaddleOCR-json\PaddleOCR-json.exe",
)

# 文本型 PDF 抽取后，若去空白字符长度小于此值，判定为扫描件，改用 OCR
SCANNED_THRESHOLD = 60

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff"}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}

# 引擎进程初始化超时（秒）。Paddle 模型加载在机械盘/首启可能较慢。
ENGINE_INIT_TIMEOUT = int(os.environ.get("UMI_OCR_INIT_TIMEOUT", "90"))


def _umi_ocr_available() -> bool:
    """本机是否存在可用的 Umi-OCR 引擎可执行文件。"""
    try:
        return bool(DEFAULT_UMI_OCR_EXE) and os.path.isfile(DEFAULT_UMI_OCR_EXE)
    except Exception:
        return False


class PaddleOCREngine:
    """调用 Umi-OCR 的 PaddleOCR-json.exe（管道模式）的薄封装。

    单例使用：首次需要 OCR 时懒启动子进程，之后复用同一进程。
    多线程并发时通过锁串行化「写指令 + 读结果」，避免管道交错。
    """

    def __init__(self, exe_path: str = None, init_timeout: int = ENGINE_INIT_TIMEOUT):
        self.exe_path = exe_path or DEFAULT_UMI_OCR_EXE
        self.init_timeout = init_timeout
        self._proc = None
        self._lock = threading.Lock()
        self._ready = False
        self._last_error = ""

    # ── 生命周期 ────────────────────────────────────────────────────────────
    def _ensure_started(self):
        """确保引擎子进程已就绪；若未启动或已崩溃则（重新）启动。"""
        if self._ready and self._proc is not None and self._proc.poll() is None:
            return
        with self._lock:
            # 双重检查，避免并发重复启动
            if self._ready and self._proc is not None and self._proc.poll() is None:
                return
            self._start()

    def _start(self):
        exe = self.exe_path
        cwd = os.path.dirname(os.path.abspath(exe))  # 必须以其所在目录为工作目录，才能找到 models/
        try:
            self._proc = subprocess.Popen(  # 管道模式
                [exe],
                cwd=cwd,
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.DEVNULL,
            )
        except Exception as e:
            self._ready = False
            self._last_error = f"启动 Umi-OCR 引擎失败：{e}"
            raise RuntimeError(self._last_error) from e

        # 用独立读取线程消费 stdout，主线程用 join(timeout) 受控等待，
        # 避免 stdout.readline() 永久阻塞（Windows 管道不支持 select 超时）。
        init_lines = []
        state = {}

        def _reader():
            try:
                while True:
                    if self._proc.poll() is not None:
                        state["dead"] = True
                        return
                    raw = self._proc.stdout.readline()
                    if not raw:
                        continue
                    line = raw.decode("utf-8", "ignore").strip()
                    if not line:
                        continue
                    init_lines.append(line)
                    if "OCR init completed." in line:
                        state["ready"] = True
                        return
            except Exception as e:
                state["err"] = str(e)

        th = threading.Thread(target=_reader, daemon=True)
        th.start()
        th.join(self.init_timeout)
        if th.is_alive():
            # 超时仍未就绪：杀掉并抛错
            try:
                self._proc.kill()
            except Exception:
                pass
            raise RuntimeError(
                f"Umi-OCR 引擎初始化超时（>{self.init_timeout}s 未就绪）。"
                f"已收到输出：{init_lines[-5:] if init_lines else '无'}"
            )
        if state.get("ready"):
            self._ready = True
            return
        if state.get("dead"):
            raise RuntimeError(
                "Umi-OCR 引擎启动后意外退出。" f"最后输出：{init_lines[-5:] if init_lines else '无'}"
            )
        raise RuntimeError(
            f"Umi-OCR 引擎初始化异常：{state.get('err', '未知')}。"
            f"最后输出：{init_lines[-5:] if init_lines else '无'}"
        )

    def stop(self):
        """关闭引擎子进程。"""
        self._ready = False
        if self._proc is not None:
            try:
                self._proc.kill()
            except Exception:
                pass
            self._proc = None

    # ── 识别接口 ────────────────────────────────────────────────────────────
    def ocr_image_bytes(self, image_bytes: bytes) -> str:
        """对一张图片的字节流做 OCR，返回拼接后的纯文本。"""
        self._ensure_started()
        with self._lock:
            if self._proc.poll() is not None:
                # 进程中途崩溃，重置后下次调用会重试
                self._ready = False
                raise RuntimeError("Umi-OCR 引擎进程已崩溃，已标记为需重启。")
            try:
                b64 = base64.b64encode(image_bytes).decode("ascii")
                cmd = json.dumps({"image_base64": b64}, ensure_ascii=True) + "\n"
                self._proc.stdin.write(cmd.encode("utf-8"))
                self._proc.stdin.flush()
            except Exception as e:
                self._ready = False
                raise RuntimeError(f"向 Umi-OCR 引擎发送指令失败：{e}") from e

            # 读取一行 JSON 结果（跳过非 JSON 的脏行）。
            # 同样用读取线程 + join(timeout)，避免 stdout.readline() 永久阻塞。
            result_state = {}
            parse_err = {}

            def _result_reader():
                try:
                    while True:
                        if self._proc.poll() is not None:
                            result_state["dead"] = True
                            return
                        raw = self._proc.stdout.readline()
                        if not raw:
                            continue
                        line = raw.decode("utf-8", "ignore").strip()
                        if not line:
                            continue
                        try:
                            res = json.loads(line)
                        except Exception:
                            continue  # 跳过非 JSON 行
                        if "code" not in res:
                            continue
                        result_state["res"] = res
                        return
                except Exception as e:
                    parse_err["err"] = str(e)

            rt = threading.Thread(target=_result_reader, daemon=True)
            rt.start()
            rt.join(self.init_timeout)
            if rt.is_alive():
                raise RuntimeError("读取 Umi-OCR 引擎结果超时。")
            if result_state.get("dead"):
                self._ready = False
                raise RuntimeError("Umi-OCR 引擎进程在识别中途退出。")
            if parse_err.get("err"):
                raise RuntimeError(f"解析 Umi-OCR 结果异常：{parse_err['err']}")
            res = result_state.get("res")
            if res is None:
                raise RuntimeError("Umi-OCR 引擎未返回有效结果。")
            if res.get("code") != 100:
                return ""  # 未识别出文字或引擎报错，按空结果处理
            data = res.get("data") or []
            return "\n".join(
                item["text"] for item in data if isinstance(item, dict) and item.get("text")
            ).strip()


# 单例（懒启动）
_paddle_engine = None
_paddle_engine_lock = threading.Lock()


def get_ocr_engine() -> PaddleOCREngine:
    """获取（或创建）Umi-OCR 引擎单例。"""
    global _paddle_engine
    if _paddle_engine is None:
        with _paddle_engine_lock:
            if _paddle_engine is None:
                _paddle_engine = PaddleOCREngine()
    return _paddle_engine


# ── EasyOCR 兜底（仅在找不到 Umi-OCR 时使用） ────────────────────────────────
_reader = None


def _get_ocr_reader():
    """懒加载 EasyOCR Reader（仅中文+英文）。"""
    global _reader
    if _reader is None:
        try:
            import easyocr
        except ImportError as e:
            raise RuntimeError(
                "未安装 EasyOCR，且未找到 Umi-OCR，无法识别扫描件/图片。"
                "请安装 Umi-OCR 或执行 pip install easyocr"
            ) from e
        _reader = easyocr.Reader(["ch_sim", "en"], gpu=False, verbose=False)
    return _reader


def _ocr_image_bytes_easyocr(image_bytes: bytes) -> str:
    """EasyOCR 兜底：先把字节存成临时图片，再识别。"""
    import tempfile

    reader = _get_ocr_reader()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as f:
        f.write(image_bytes)
        tmp = f.name
    try:
        results = reader.readtext(tmp, detail=0, paragraph=True)
    finally:
        try:
            os.unlink(tmp)
        except Exception:
            pass
    return "\n".join(r.strip() for r in results if r and r.strip())


def _ocr_image_bytes(image_bytes: bytes) -> str:
    """统一 OCR 入口：优先 Umi-OCR，失败或不可用时回退 EasyOCR。"""
    if _umi_ocr_available():
        try:
            return get_ocr_engine().ocr_image_bytes(image_bytes)
        except Exception as e:
            # Umi-OCR 调用异常，退一步用 EasyOCR，不让单点故障阻断整批
            print(f"[warn] Umi-OCR 识别失败，回退 EasyOCR：{e}", file=sys.stderr)
    return _ocr_image_bytes_easyocr(image_bytes)


# ── 各格式提取 ──────────────────────────────────────────────────────────────
def _extract_pdf_text(path: str) -> str:
    """尝试从 PDF 抽取文本；文本不足则渲染图片 OCR。返回纯文本。"""
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    try:
        texts = []
        for page in doc:
            txt = page.get_text().strip()
            if txt:
                texts.append(txt)
        full = "\n".join(texts).strip()
        if len(full) >= SCANNED_THRESHOLD:
            return full
        # 文本不足 → 扫描件，逐页渲染图片 OCR
        ocr_parts = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            ocr_parts.append(_ocr_image_bytes(pix.tobytes("png")))
        return "\n".join(p for p in ocr_parts if p.strip())
    finally:
        doc.close()


def _extract_docx_text(path: str) -> str:
    import docx

    document = docx.Document(path)
    paras = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paras.append(cell.text.strip())
    return "\n".join(paras).strip()


def extract_text(file_path: str, filename: str | None = None) -> str:
    """统一入口：根据扩展名选择提取方式，返回纯文本。

    file_path: 本地的临时文件路径
    filename:  原始文件名（用于推断扩展名，可选）
    """
    ext = (filename or file_path).lower()
    ext = Path(ext).suffix

    if ext in _DOCX_EXTS:
        return _extract_docx_text(file_path)
    if ext in _PDF_EXTS:
        return _extract_pdf_text(file_path)
    if ext in _IMAGE_EXTS:
        with open(file_path, "rb") as f:
            return _ocr_image_bytes(f.read())

    # 未知类型：尝试按 PDF 处理，失败再当图片 OCR，再失败报错
    try:
        return _extract_pdf_text(file_path)
    except Exception:
        pass
    try:
        with open(file_path, "rb") as f:
            return _ocr_image_bytes(f.read())
    except Exception as e:
        raise ValueError(f"不支持的文件类型或解析失败：{ext or '未知'}（{e}）")


def ocr_backend_info() -> dict:
    """返回当前 OCR 后端信息，供前端展示。"""
    if _umi_ocr_available():
        return {
            "backend": "umi-ocr",
            "name": "Umi-OCR (PaddleOCR-json)",
            "exe": DEFAULT_UMI_OCR_EXE,
            "note": "调用本机已安装的 Umi-OCR 引擎（PaddleOCR），中文识别精度高。",
        }
    return {
        "backend": "easyocr",
        "name": "EasyOCR",
        "exe": None,
        "note": "未找到 Umi-OCR 引擎，已回退到 EasyOCR。",
    }


# 进程退出时关闭引擎子进程
atexit.register(lambda: _paddle_engine.stop() if _paddle_engine else None)
